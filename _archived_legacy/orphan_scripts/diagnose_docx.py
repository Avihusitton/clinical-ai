#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
diagnose_docx.py -- כלי אבחון: איפה בדיוק חי כל טקסט בתוך קובץ docx.

בלי --search: מדפיס מלאי מבני (סעיפים, כותרות עליונות/תחתונות, יחס
פסקאות-כותרת, טבלאות, תיבות טקסט) - שימושי לפני שמריצים את הצינור המלא,
כדי לדעת למה לצפות.

עם --search "טקסט": מאתר בדיוק היכן טקסט מסוים מופיע - body, header,
footer, table, או textbox/shape (שהצינור הרגיל לא קורא בכלל).

הרצה:
    python diagnose_docx.py path/to/file.docx
    python diagnose_docx.py path/to/file.docx --search "12.3.24"
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn


def inventory(path: Path) -> None:
    document = docx.Document(str(path))
    print(f"=== מלאי מבני: {path.name} ===\n")
    print(f"סעיפים (sections): {len(document.sections)}")
    for i, s in enumerate(document.sections):
        htext = " / ".join(p.text.strip() for p in s.header.paragraphs if p.text.strip())
        ftext = " / ".join(p.text.strip() for p in s.footer.paragraphs if p.text.strip())
        print(f"  סעיף {i}: כותרת עליונה: {htext or '(ריקה)'} | כותרת תחתונה: {ftext or '(ריקה)'}")

    non_empty = [p for p in document.paragraphs if p.text.strip()]
    headings = [p for p in non_empty
                if (p.style.name or "").lower().startswith("heading") or "כותרת" in (p.style.name or "")]
    ratio = 100 * len(headings) / max(len(non_empty), 1)
    print(f"\nפסקאות גוף (לא ריקות): {len(non_empty)}")
    print(f"מתוכן בסגנון כותרת: {len(headings)} ({ratio:.0f}%)")
    print(f"טבלאות: {len(document.tables)}")

    textbox_count = len(document.element.body.findall(".//" + qn("w:txbxContent")))
    print(f"תיבות טקסט (textboxes) שזוהו ב-XML גולמי: {textbox_count}")
    if textbox_count:
        print("  ⚠️  יש תיבות טקסט - התוכן שם לא נקרא כלל על ידי ingestion_pipeline.py!")

    print("\nרמז לסיווג DocumentTypeClassifier:")
    first_header = " / ".join(p.text.strip() for p in document.sections[0].header.paragraphs if p.text.strip())
    print(f"  כותרת עמוד ראשונה קצרה (<=80 תווים)? {'כן' if first_header and len(first_header) <= 80 else 'לא'}")
    print(f"  יחס כותרות >= 15%? {'כן' if ratio >= 15 else 'לא'}")


def find_text(path: Path, needle: str) -> None:
    document = docx.Document(str(path))
    print(f"\n=== חיפוש '{needle}' ב-{path.name} ===\n")
    found = False

    for i, p in enumerate(document.paragraphs):
        if needle in p.text:
            print(f"[body] פסקה {i}: ...{p.text.strip()[:150]}...")
            found = True

    for si, s in enumerate(document.sections):
        for p in s.header.paragraphs:
            if needle in p.text:
                print(f"[header, סעיף {si}]: {p.text.strip()}")
                found = True
        for p in s.footer.paragraphs:
            if needle in p.text:
                print(f"[footer, סעיף {si}]: {p.text.strip()}")
                found = True

    for ti, table in enumerate(document.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if needle in cell.text:
                    print(f"[table {ti}, שורה {ri}, עמודה {ci}]: {cell.text.strip()[:150]}")
                    found = True

    for txbx in document.element.body.findall(".//" + qn("w:txbxContent")):
        txt = "".join(txbx.itertext())
        if needle in txt:
            print(f"[textbox/shape - XML גולמי, לא נקרא ע\"י הצינור הרגיל]: {txt.strip()[:150]}")
            found = True

    if not found:
        print("לא נמצא באף אחד מהמקומות שנבדקו (body/header/footer/table/textbox).")


def main() -> None:
    ap = argparse.ArgumentParser(description="אבחון מבנה קובץ docx")
    ap.add_argument("docx_path", type=Path)
    ap.add_argument("--search", default=None, help="טקסט לחיפוש מדויק בתוך הקובץ")
    args = ap.parse_args()

    if not args.docx_path.exists():
        sys.exit(f"קובץ לא נמצא: {args.docx_path}")

    inventory(args.docx_path)
    if args.search:
        find_text(args.docx_path, args.search)


if __name__ == "__main__":
    main()
